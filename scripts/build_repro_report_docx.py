from __future__ import annotations

import csv
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "outputs" / "report_assets"
OUT = ROOT / "outputs" / "Nature613_复现报告_按参考格式重排.docx"


CN_BODY = "仿宋"
CN_TITLE = "楷体"
CN_HEADING = "黑体"
EN_FONT = "Times New Roman"


def set_font(run, east_asia=CN_BODY, ascii_font=EN_FONT, size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_common(p, first_line=True, after=0):
    fmt = p.paragraph_format
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    fmt.line_spacing = 1.25
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="000000", width="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), width)
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Nature 613 论文复现报告")
    set_font(r, east_asia=CN_BODY, size=9)


def paragraph(doc, text="", first_line=True, bold=False, font=CN_BODY, size=12, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_common(p, first_line=first_line, after=0)
    r = p.add_run(text)
    set_font(r, east_asia=font, size=size, bold=bold)
    return p


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_font(r, east_asia=CN_TITLE, size=22, bold=True)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, east_asia=CN_HEADING, size=14, bold=True)


def h2(doc, text, indented=True):
    p = doc.add_paragraph()
    if indented:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, east_asia=CN_HEADING, size=12, bold=True)


def add_note_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="777777", width="4")
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.width = Cm(widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 or row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, east_asia=CN_HEADING if row_idx == 0 else CN_BODY, size=10.5, bold=(row_idx == 0 or col_idx == 0))
    doc.add_paragraph()
    return table


def add_figure_table(doc, image_name, caption, width_cm=14.0):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="B7B7B7", width="3")
    for row in table.rows:
        row.cells[0].width = Cm(14.0)
        set_cell_margins(row.cells[0], top=80, bottom=80, start=80, end=80)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(ASSETS / image_name), width=Cm(width_cm))

    ccell = table.cell(1, 0)
    ccell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = ccell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(caption)
    set_font(r, east_asia=CN_BODY, size=10)
    doc.add_paragraph()


def parse_test_metrics():
    text = (ASSETS / "retrained_test_eval.log").read_text(encoding="utf-8", errors="ignore")
    metrics = {}
    for key in ["samples", "top1_accuracy", "top3_accuracy", "grouped_99_accuracy", "mean_group_size_99"]:
        match = re.search(rf"{key}:\s+([0-9.]+)", text)
        if match:
            value = match.group(1)
            metrics[key] = int(value) if key == "samples" else float(value)
    return metrics


def load_last_epochs(n=8):
    rows = []
    with (ASSETS / "training_metrics.csv").open("r", encoding="utf-8", newline="") as f:
        for row in csv.DictReader(f):
            rows.append(row)
    return rows[-n:]


def main():
    summary = json.loads((ASSETS / "training_summary.json").read_text(encoding="utf-8"))
    low_classes = json.loads((ASSETS / "lowest_per_class_accuracy.json").read_text(encoding="utf-8"))
    test = parse_test_metrics()
    last_epochs = load_last_epochs()

    doc = Document()
    configure_document(doc)

    title(doc, "Nature 613 论文模型复现报告")
    paragraph(doc, "——基于官方代码的训练过程、模型产物与测试结果整理", first_line=False, font=CN_HEADING, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}", first_line=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    h1(doc, "一、复现任务概述")
    paragraph(doc, "本次复现围绕 Nature 613 论文相关官方模型代码展开，目标是在本机 GPU 环境中完成从训练启动、训练日志保存、模型权重生成到独立测试集评估的完整流程。复现对象为二十类有机反应机理分类模型，即 M1 至 M20 的多分类任务。")
    paragraph(doc, "从训练日志看，模型已经完成完整训练流程。训练原计划为 3000 个 epoch，实际在第 1931 个 epoch 触发 early stopping，并恢复至验证集表现最优的第 1631 个 epoch 权重。最终模型文件保存为 M1_20_model.keras，说明训练过程已经成功产出可复核模型。")
    paragraph(doc, "与只进行冒烟训练不同，本次结果来自长时间正式训练日志。日志中记录的平均单 epoch 时间约为 112.34 秒，按已记录 epoch 估算总训练时长约 60.26 小时，能够反映模型在本机 GPU 条件下完成了较充分的收敛过程。")

    h1(doc, "二、运行环境与复现材料")
    paragraph(doc, "训练运行在 WSL Ubuntu-24.04 环境中，使用 TensorFlow 调用本机 NVIDIA GeForce RTX 4060 Laptop GPU。日志显示 TensorFlow 成功创建 GPU 设备，显存分配约 5563 MB，并加载 cuDNN，说明训练主体确实运行在 GPU 环境下。")
    add_note_table(
        doc,
        [
            ["项目", "内容"],
            ["复现目录", "D:\\复现\\nature613_repro"],
            ["训练日志", "outputs/full_train.log"],
            ["保存模型", "M1_20_model.keras"],
            ["测试概率文件", "outputs/report_assets/retrained_test_probabilities.npy"],
            ["测试标签文件", "outputs/report_assets/retrained_test_labels.npy"],
            ["报告配图目录", "outputs/report_assets/*.png"],
        ],
        [4.2, 9.5],
    )
    paragraph(doc, "本次报告采用训练日志、测试评估输出、训练曲线、损失曲线、验证集趋势图和混淆矩阵作为主要证据材料。上述材料均来自复现过程中的实际输出，而不是手工估算。")

    h1(doc, "三、训练过程分析")
    h2(doc, "1. 训练终止与最佳权重")
    paragraph(doc, f"训练共记录 {summary['epochs_recorded']} 个 epoch。最后一个完整记录为 epoch {summary['final_epoch']}，其训练集 categorical accuracy 为 {summary['final_train_accuracy']:.4f}，训练损失为 {summary['final_train_loss']:.4f}；验证集 categorical accuracy 为 {summary['final_val_accuracy']:.4f}，验证损失为 {summary['final_val_loss']:.4f}。")
    paragraph(doc, f"验证集最优点出现在 epoch {summary['best_val_accuracy_epoch']}。该 epoch 的 val_categorical_accuracy 为 {summary['best_val_accuracy']:.4f}，val_loss 为 {summary['best_val_loss']:.4f}。训练结束时日志明确显示 early stopping，并将模型权重恢复到第 {summary['best_val_accuracy_epoch']} 个 epoch。")
    add_note_table(
        doc,
        [
            ["指标", "数值"],
            ["记录 epoch 数", summary["epochs_recorded"]],
            ["最终 epoch", summary["final_epoch"]],
            ["最终训练准确率", f"{summary['final_train_accuracy']:.4f}"],
            ["最终验证准确率", f"{summary['final_val_accuracy']:.4f}"],
            ["最佳验证 epoch", summary["best_val_accuracy_epoch"]],
            ["最佳验证准确率", f"{summary['best_val_accuracy']:.4f}"],
            ["最佳验证损失", f"{summary['best_val_loss']:.4f}"],
        ],
        [5.0, 8.7],
    )

    h2(doc, "2. 收敛曲线观察")
    paragraph(doc, "从准确率曲线看，模型在训练早期快速提升，此后进入缓慢爬升和小幅波动阶段。验证集准确率在后半段长期稳定在 0.80 以上，说明模型已经学到主要类别区分特征，但仍存在部分类别的局部混淆。")
    add_figure_table(doc, "accuracy_curve.png", "图1 训练集与验证集 categorical accuracy 随 epoch 的变化。", width_cm=13.5)
    paragraph(doc, "从损失曲线看，训练损失和验证损失在早期快速下降，随后进入较平稳的收敛区间。验证损失后期仍有少量尖峰，说明不同 epoch 下模型对某些验证样本的判别边界仍会发生波动，这也是 early stopping 选择最佳权重的必要原因。")
    add_figure_table(doc, "loss_curve.png", "图2 训练集与验证集 categorical cross-entropy loss 随 epoch 的变化。", width_cm=13.5)
    paragraph(doc, "单独观察验证准确率趋势可以看到，最佳 checkpoint 位于 epoch 1631，后续训练虽然仍有局部波动，但没有持续产生更好的验证集综合表现。因此恢复第 1631 个 epoch 的权重是合理的。")
    add_figure_table(doc, "validation_accuracy_trend.png", "图3 验证准确率趋势与最佳 checkpoint 标记。", width_cm=13.5)

    h2(doc, "3. 训练后期日志抽查")
    rows = [["Epoch", "train acc", "train loss", "val acc", "val loss"]]
    for r in last_epochs:
        rows.append([
            str(int(float(r["epoch"]))),
            f"{float(r['categorical_accuracy']):.4f}",
            f"{float(r['loss']):.4f}",
            f"{float(r['val_categorical_accuracy']):.4f}",
            f"{float(r['val_loss']):.4f}",
        ])
    add_note_table(doc, rows, [2.2, 2.6, 2.6, 2.6, 2.6])
    paragraph(doc, "训练最后若干个 epoch 的训练集准确率与验证集准确率接近，二者没有出现明显大幅背离，说明此时模型主要处于收敛后的细微波动阶段，而不是严重过拟合后继续恶化。")

    h1(doc, "四、测试集评估结果")
    paragraph(doc, f"使用保存后的 M1_20_model.keras 对官方 test_subset 中 20 timepoints、1% noise 分支进行预测，测试样本量为 {test['samples']:,}。测试结果显示，模型 Top-1 accuracy 为 {test['top1_accuracy']:.6f}，Top-3 accuracy 为 {test['top3_accuracy']:.6f}，99% 概率分组准确率为 {test['grouped_99_accuracy']:.6f}，对应平均分组大小为 {test['mean_group_size_99']:.6f}。")
    add_note_table(
        doc,
        [
            ["评价指标", "结果"],
            ["测试样本数", f"{test['samples']:,}"],
            ["Top-1 accuracy", f"{test['top1_accuracy']:.6f}"],
            ["Top-3 accuracy", f"{test['top3_accuracy']:.6f}"],
            ["99% 概率分组准确率", f"{test['grouped_99_accuracy']:.6f}"],
            ["99% 分组平均大小", f"{test['mean_group_size_99']:.6f}"],
        ],
        [5.0, 8.7],
    )
    paragraph(doc, "Top-1 accuracy 代表模型单一最大概率类别的直接命中率；Top-3 accuracy 说明真实类别是否落在概率最高的前三个候选中；99% 概率分组准确率则更接近化学机理检索或候选机理推荐场景，即模型给出一个高置信候选集合，再由后续规则或人工判断进一步筛选。")
    add_figure_table(doc, "confusion_matrix.png", "图4 100,000 个测试样本上的类别归一化混淆矩阵。", width_cm=12.6)

    h1(doc, "五、误差集中类别与结果解释")
    paragraph(doc, "从混淆矩阵和逐类准确率看，大部分机理类别能够被模型稳定识别，主对角线颜色较深，说明预测结果总体集中在正确类别上。但 M11 与 M14 的 Top-1 accuracy 明显低于其他类别，是本次复现模型的主要误差来源。")
    rows = [["类别", "样本数", "正确数", "Top-1 accuracy"]]
    for item in low_classes:
        rows.append([item["class"], item["samples"], item["correct"], f"{item['accuracy']:.4f}"])
    add_note_table(doc, rows, [2.6, 3.4, 3.4, 4.3])
    paragraph(doc, "M11 与 M14 的准确率分别为 0.5366 和 0.5852，说明这两类与其他机理之间的特征边界更接近，或者在当前输入条件下能够被模型利用的差异特征不足。M8、M15 与 M20 的准确率也低于整体平均水平，但仍明显高于 M11 与 M14。")
    paragraph(doc, "从复现角度看，这种误差分布并不代表训练失败，而是提示后续若要进一步逼近论文报告值，需要重点检查类别定义、数据预处理分支、随机种子、官方预训练权重差异、测试噪声设置以及 timepoints 选择对这些低准确率类别的影响。")

    h1(doc, "六、与原论文和官方模型结果的比较")
    paragraph(doc, "原论文为 Burés 与 Larrosa 发表在 Nature 613 的 Organic reaction mechanism classification using machine learning。论文的核心结论是：深度神经网络可以从普通动力学数据中自动识别有机反应机理类别，并且在少量时间点或含噪声条件下仍保持较好的分类能力。")
    paragraph(doc, "为了避免把不同测试条件混为一谈，本报告将比较分成两个层次。第一层是同一测试分支比较，即本次重训模型与官方预训练模型都在 standard_tp20_noise1，也就是 20 timepoints、1% noise 分支上评估。第二层是与论文核心标准分支比较，即官方预训练模型在 tp6_noise0 分支上的表现；该分支在本地复现文档中被标记为论文核心表述对应的标准测试分支。")
    h2(doc, "1. 同一测试分支上的比较")
    paragraph(doc, "在 standard_tp20_noise1 分支上，本次重训模型的 Top-1、Top-3 与 99% 分组准确率均与官方预训练模型非常接近。其中 Top-1 accuracy 从官方预训练模型的 0.891460 变为本次重训模型的 0.892090，差值为 +0.000630；Top-3 accuracy 差值为 +0.000150；99% 分组准确率差值为 +0.000040。")
    add_note_table(
        doc,
        [
            ["指标", "官方预训练模型", "本次重训模型", "差值"],
            ["Top-1 accuracy", "0.891460", "0.892090", "+0.000630"],
            ["Top-3 accuracy", "0.995430", "0.995580", "+0.000150"],
            ["99% 分组准确率", "0.998350", "0.998390", "+0.000040"],
            ["99% 平均分组大小", "1.673070", "1.678870", "+0.005800"],
        ],
        [3.5, 3.5, 3.5, 3.2],
    )
    paragraph(doc, "这一结果说明：在相同测试分支下，本次完整训练得到的模型与官方预训练模型基本对齐，差异处于很小范围内。由于训练随机初始化、TensorFlow/Keras 版本、GPU 数值路径、early stopping checkpoint 和数据读取顺序可能存在差异，完全逐位一致并不现实；但从指标量级看，本次重训结果可以认为复现到了官方模型在该分支上的性能水平。")
    h2(doc, "2. 与论文核心标准分支的比较")
    paragraph(doc, "本地复现记录显示，论文核心表述对应的标准测试分支更接近 tp6_noise0，即 6 个时间点、0% 噪声条件。官方预训练模型在该分支上的 Top-1 accuracy 为 0.926390，Top-3 accuracy 为 1.000000，99% 分组准确率为 0.999620。")
    add_note_table(
        doc,
        [
            ["对比对象", "测试分支", "Top-1", "Top-3", "99% 分组"],
            ["论文/官方核心结果", "tp6_noise0", "0.926390", "1.000000", "0.999620"],
            ["本次重训评估结果", "standard_tp20_noise1", "0.892090", "0.995580", "0.998390"],
            ["表面差值", "不同分支", "-0.034300", "-0.004420", "-0.001230"],
        ],
        [3.4, 4.0, 2.1, 2.1, 2.1],
    )
    paragraph(doc, "需要强调的是，上表最后一行只能作为总体参照，不能作为严格同条件差距，因为两个结果来自不同测试分支。本次报告目前的重训模型测试是在 20 timepoints、1% noise 条件下完成，而论文核心标准结果对应的是 6 timepoints、0% noise 条件。若要做最严格的论文数值复现，应继续使用本次保存的 M1_20_model.keras 在 tp6_noise0 分支上重新评估，再与 0.926390、1.000000 和 0.999620 直接比较。")
    paragraph(doc, "从现有证据看，本次重训模型在同分支上已经贴近官方预训练模型；与论文核心分支的表面差距主要来自测试分支不同，而不是能够直接说明训练失败。后续最有价值的补充实验，是对同一个重训模型依次运行 tp6_noise0、tp6_noise1、tp6_noise5 和 tp2_noise1，形成完整的噪声鲁棒性与时间点数量对照表。")

    h1(doc, "七、与论文复现要求的对应关系")
    paragraph(doc, "本次复现至少完成了三个关键环节。第一，正式训练已经完成，而不是仅完成冒烟测试；第二，训练结束后生成了完整模型文件，并且该模型可以被重新加载用于测试集预测；第三，测试评估产物已经保存为概率数组、标签数组、混淆矩阵和指标文件，具备后续复查条件。")
    paragraph(doc, "从工程可复核性看，完整训练日志 outputs/full_train.log 是最重要证据。该日志不仅记录了训练过程，也记录了 GPU 设备创建、batch size、每个 epoch 的训练指标、验证指标、early stopping 和模型保存信息。")
    paragraph(doc, "从结果可信度看，测试集评估并非只报告训练集指标，而是使用保存模型对独立测试样本重新推理。测试集 Top-1 accuracy 达到 0.892090，Top-3 accuracy 达到 0.995580，说明复现模型已经具备较强的类别排序能力。")

    doc.add_page_break()
    h1(doc, "八、可复核产物清单")
    paragraph(doc, "为便于后续继续核验、补实验或写入更完整论文复现材料，本次复现过程保留了如下产物：")
    add_note_table(
        doc,
        [
            ["产物", "路径或说明"],
            ["完整训练日志", "outputs/full_train.log"],
            ["训练指标 CSV", "outputs/report_assets/training_metrics.csv"],
            ["训练摘要 JSON", "outputs/report_assets/training_summary.json"],
            ["保存模型", "M1_20_model.keras"],
            ["测试预测概率", "outputs/report_assets/retrained_test_probabilities.npy"],
            ["测试真实标签", "outputs/report_assets/retrained_test_labels.npy"],
            ["混淆矩阵数据", "outputs/report_assets/confusion_matrix.csv"],
            ["报告配图", "accuracy_curve.png、loss_curve.png、validation_accuracy_trend.png、confusion_matrix.png"],
        ],
        [4.1, 9.6],
    )

    h1(doc, "九、复现结论")
    paragraph(doc, "本次 Nature 613 模型复现已经完成正式训练、模型保存和独立测试评估。训练在 epoch 1931 触发 early stopping，并恢复 epoch 1631 的最佳验证集权重；保存模型在 100,000 个测试样本上取得 Top-1 accuracy 0.892090、Top-3 accuracy 0.995580。")
    paragraph(doc, "与官方预训练模型在同一 standard_tp20_noise1 分支上的结果相比，本次重训模型 Top-1 accuracy 高出 0.000630，Top-3 accuracy 高出 0.000150，99% 分组准确率高出 0.000040，说明本次训练结果已经基本复现官方模型在该测试分支上的表现。")
    paragraph(doc, "综合训练曲线、验证集趋势、测试集混淆矩阵和与官方结果的对照，本次复现说明本机 GPU 环境能够稳定完成该模型训练，生成的模型具有可用的分类能力。后续若要进一步严格对齐论文核心标准分支，应使用保存模型补测 tp6_noise0，并优先围绕低准确率类别 M11、M14 进行误差分析。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
